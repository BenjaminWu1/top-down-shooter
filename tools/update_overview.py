#!/usr/bin/env python3
"""
Sync Game-Overview.docx with the current game (index.html).

A small, repeatable maintenance tool. It opens the Word doc with python-docx and
applies a curated list of TARGETED edits:

  * REPLACEMENTS - rewrite a whole paragraph, located by a unique substring of
    its current text (paragraph STYLE is preserved; the text is replaced).
  * INSERTS      - add new bullet paragraphs immediately before a heading.

It is IDEMPOTENT, so re-running after future game changes is safe:
  * a replacement whose "find" text is already gone is reported "already current"
    and skipped (pick "find" substrings that do NOT appear in the new text);
  * an insert whose leading text already exists in the doc is not duplicated.

When the game changes again, DON'T hand-edit the docx - add an entry here and
re-run:  python tools/update_overview.py

Requires:  pip install python-docx     (NOT stdlib - unlike tools/validate.py)
"""
import os
import sys

try:
    import docx
except ImportError:
    print("ERROR: python-docx is not installed. Run:  pip install python-docx",
          file=sys.stderr)
    sys.exit(2)

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.normpath(os.path.join(HERE, "..", "Game-Overview.docx"))

# ---------------------------------------------------------------------------
# Edits. Each REPLACEMENTS entry is (find_substring, new_full_paragraph_text).
# The find_substring must uniquely identify ONE paragraph and must NOT occur in
# the new text (so a second run is a no-op).
# ---------------------------------------------------------------------------
REPLACEMENTS = [
    # 1. The big picture - refresh the line count.
    ("index.html (~5,963 lines)",
     "The entire game is one file: index.html (~6,224 lines). One <script> block, "
     "vanilla JS, no build, no dependencies, no external assets. Sprites are inline "
     "ASCII-art arrays; all SFX are synthesized via Web Audio. It deploys by pushing "
     "to main (GitHub Pages serves the raw file)."),

    # 2. Canvas/world/camera - 4 tiers + MAP_EXPAND, camera now on every tier.
    ("Three world-size tiers",
     "Four world-size tiers, chosen in startLevel(idx) via worldScale = floor(idx/10)+1 "
     "and then enlarged by a single constant MAP_EXPAND = 1.15 (about 1.32x area), "
     "rounded: L1-L10 = 552x311, L11-L20 = 1104x621, L21-L30 = 1656x932, "
     "L31-L40 = 2208x1242. Arena AREA grows with the block. Because even the smallest "
     "tier now exceeds the 480x270 viewport, the camera follows the player on EVERY "
     "tier (updateCamera collapses to 0,0 only when there is no room to scroll). Each "
     "10-level block also gets a distinct theme: forest -> lava -> ice -> void."),

    # 7. Weapons - fireWeapon no longer a form-changing cascade; enhancers only.
    ("fixed priority cascade: rocket",
     "fireWeapon(mx,my) ALWAYS fires the class basic attack - powerups NEVER change its "
     "form, they only ENHANCE it. Four timed (~8s) enhancers stack: rapidTime halves "
     "cooldown, damageTime doubles damage (with damageMult), pierceTime sets the bullet "
     "pierce flag (shots punch through enemies), multiTime fires a parallel DOUBLE shot "
     "(two bullets offset +/-3 perpendicular to aim, same projectile). Class defaults: "
     "Scout blue-dot SMG (dmg 0.85, cd 0.06), Tank slug (dmg 2, radius 3), Soldier "
     "pistol (dmg 1)."),

    # 7. Bullet flags - plasma combo gone; homing/bounces are enemy-only now.
    ("Plasma = pierce + bounces",
     "Bullet flags (createBullet): pierce (uses hitSet, does not die on hit), explosive "
     "(explode()), homing (enemy bullets steer toward the player), bounces (reflects off "
     "edges, clears hitSet), isGrenade, fromAlly. The homing/bounces flags are now used "
     "only by ENEMY AI - the old player rocket/laser/plasma weapons were removed and "
     "their bullet sprites are orphaned."),

    # 7. Pickups - new enhancer set, removed the 6 form-changers.
    ("Pickups (applyPickup): equipment - health, spread",
     "Pickups (applyPickup): WEAPON ENHANCERS (timed ~8s, never change the attack's "
     "form) - rapid (half cooldown), pierce (shots punch through), multi (single -> "
     "double shot), damage (2x); defenses - health, shield, armor; utility powerups - "
     "speed, slowmo, score2x, magnet; class resources - fuel, beverage, battery. The 6 "
     "old form-changing weapons (rocket/laser/plasma/minigun/spread/homing) were removed "
     "as pickups. Three drop pools in killEnemy (Tank->fuel, Soldier->beverage, "
     "Scout->battery), each MUST sum to 1.0 (validator C5). 22% drop chance per non-boss "
     "kill (55% for elite mini-bosses). Bosses give a guaranteed enhancer/defense drop "
     "plus a health."),

    # 9. Bosses - updateTwin/updateBomber + the new warden/harrier are medium/trans bosses.
    ("they are the L31-L40 MID-BOSSES",
     "updateTwin and updateBomber are no longer orphaned - they are reused as "
     "medium/transitional bosses in the staged L11-40 encounters, alongside two NEW "
     "boss AIs (updateWarden, updateHarrier) used ONLY as medium/transitional bosses. "
     "Each L11-40 level carries midBoss (+ midBoss2 from L21, + transBoss on L31-40) "
     "naming the medium/transitional kinds, each a DIFFERENT kind than its end boss."),

    # 11. Draw/HUD - enemy counter removed; enhancer chips above the RMB bar.
    ("Boss HP bar, enemy counter, arrival flash are all screen-space",
     "draw() flow on game states: save -> translate(-camX,-camY) -> drawBackground -> "
     "drawGame() (entities + particles) -> restore, then un-translated drawHUD() + "
     "overlays. The boss HP bar and boss-arrival flash are screen-space. (The on-screen "
     "ENEMIES-left counter was removed.) The four basic-attack enhancer chips "
     "(RAPID/PIERCE/MULTI/DMG) render top-left, stacked ABOVE the RMB resource bar "
     "(FUEL/ENERGY/POWER); movement/utility buffs sit in the bottom-left strip."),

    # 12. Validation - mention the headless runtime+balance check too. (The "find"
    # phrase must NOT recur in the new text, or this re-fires every run.)
    ("asserts ~18 balance properties",
     "tools/validate.py (pure stdlib) is a static invariant linter (not runtime): one "
     "<script>, brace balance, 40 levels, every mix sums to 1.0 (C4), all 3 drop pools "
     "sum to 1.0 (C5), boss tiers, skill wiring, GAMEPLAY_KEYS coverage. Run "
     "python tools/validate.py after editing LEVELS, drop pools, skills, or keys; CI "
     "runs it on push as a signal only (it never blocks the Pages deploy). There is "
     "ALSO an optional (non-CI) runtime + balance check, tools/headless_check.py (needs "
     "pip install py_mini_racer): it executes the game in V8 with a stubbed canvas/DOM "
     "across all 40 levels and checks ~23 balance properties - including the STAGED "
     "boss flow (1/2/3 mediums by block, the 10-15s surge gap, and the L31-40 0s "
     "transitional-to-final handoff)."),

    # "Things to keep in mind" - the cascade reminder is obsolete; enhancers only.
    ("The weapon priority cascade in fireWeapon",
     "Weapon pickups only ENHANCE the basic attack (rapid/pierce/multi/damage) - they "
     "must NEVER change its projectile form."),

    # 3. STATE machine - SHOP state added.
    ("CHARACTER_SELECT, ASSISTANT_SELECT. Both update() and draw() dispatch",
     "STATE enum: MENU, HOWTO, PLAYING, LEVEL_COMPLETE, GAME_OVER, VICTORY, "
     "CHARACTER_SELECT, ASSISTANT_SELECT, SHOP. Both update() and draw() dispatch on "
     "the state variable."),

    # 3. Run setup flow - now a loadout screen + shop, not 'pick exactly 2'.
    # (SHOP button moved lower-left -> top-right under PARAMETERS; re-key on the OLD text.)
    ("opens from a lower-left",
     "Run setup flow: MENU (pick level) -> CHARACTER_SELECT (pick class) -> "
     "ASSISTANT_SELECT (the LOADOUT screen: equip 0-2 OWNED assistants + assign "
     "owned skills to X/C/B/V) -> PLAYING. SHOP (the ARMORY) opens from a top-right "
     "MENU button, stacked under the PARAMETERS button. The real entry point is "
     "startSelectedRun() (uses selectedChar + selectedLevel + the persistent profile)."),

    # Shop now buys AND upgrades assistants (10 levels); UI layout move.
    # (find anchors on current docx text and must NOT recur in new, or it re-fires.)
    ("and extra assistants; purchases persist",
     "SHOP (STATE.SHOP, the ARMORY): spend Gold on tiered skills (cheap/small -> "
     "expensive/big) and on assistants - which you BUY then UPGRADE through 10 levels "
     "(the gold cost rises each level); purchases persist. LOADOUT screen: F is locked "
     "to HEAL, X/C/B/V unlock +1 per 5 levels and are filled from owned skills. The "
     "GUIDE button now sits top-LEFT; the persistent LV/XP bar shows in the top-RIGHT "
     "corner with the COINS total (a specific number) on its own line just below the "
     "XP bar."),

    # 5. Player classes - active skills are now a loadout, not the fixed kit.
    # (find must NOT recur in the new text, or it re-fires every run.)
    ("off them. Display names and live stats:",
     "Three classes; class keys never change (soldier/scout/tank) so all logic keys "
     "off them. Each keeps a class-locked LMB + held-RMB; the active X/C/B/V skills "
     "are now a PROFILE LOADOUT (F=HEAL locked, +1 slot per 5 levels) drawn from the "
     "SKILLS catalog, NOT the per-class CHARACTERS.skills kit (now vestigial / kept "
     "only for validator C7). Display names and live stats:"),

    # 10. Allies - owned + equip 0..2 + per-ally 10-level gold upgrades.
    # (find anchors on current docx text and must NOT recur in new, or it re-fires.)
    ("the old exactly-2 rule is gone",
     "Roster in ASSISTANTS (drone, henchman, nunchaku, bomber, poison); the player "
     "OWNS a default drone + any bought in the SHOP, and EQUIPS 0-2 of them on the "
     "loadout screen (the equipped set is default-removable). Each ally also has a "
     "per-account UPGRADE LEVEL 1..10 (profile.assistantLevels, raised with gold in "
     "the SHOP): createAssistant runs applyAssistantLevel, treating the per-kind stats "
     "as the level-10 reference and scaling HP/damage/range/attack-speed down by "
     "assistLevelFrac (~0.35 at L1 -> 1.0 at L10), so a fresh ally is much weaker and "
     "upgrades buy power back. Equipped allies spawn at run start "
     "(spawnSelectedAssistants) and rebuild each level (refreshAssistants)."),

    # Difficulty curve bullet -> FOUR-PHASE + the L31-40 spike + the two new enemies.
    # (Rewrites the bullet inserted earlier; text MUST match the INSERTS version so the
    # insert step then sees it present and skips it. find anchors on the OLD bullet text.)
    ("THREE-PHASE difficulty curve (difficultyRamp",
     "FOUR-PHASE difficulty curve (difficultyRamp / enemyScale / bossHpScale / "
     "bossSpdScale): scaling is gentle through L1-15, moderate L16-24, sharp L25-30, then a "
     "STEEP L31-40 spike so the late game depends on meta-progression upgrades. The enemy "
     "HP/speed ramp's per-level step is capped at 0.2, so most of the L31-40 spike is carried "
     "by boss HP/speed plus enemyDmgScale(idx) (1.0 through L30, up to ~2.2 at L40, applied to "
     "enemy contact/bullet/gas damage). Two new regular enemies join the L11-40 spawn mixes: "
     "weaver (evasive bullet-dodger) and spitter (drops telegraphed poison pools). DROP "
     "variety is also level-gated in killEnemy - basic & sparse in L1-15 "
     "(health/resource/rapid/shield), widening at L16 (pierce/multi/damage/speed/magnet) and "
     "L25 (full pool) - while the static drop arrays still sum to 1.0 (validator C5) and the "
     "runtime filter renormalizes."),

    # Mid-boss bullet -> the full STAGED encounter structure. (Same rewrite-in-place trick.)
    ("L31-L40 mid-boss + gap: a mid-boss",
     "STAGED boss encounters by 10-level block (bossPhase state machine): L1-10 end "
     "boss only; L11-20 one medium boss at 50% progress, then a surge gap, then the end "
     "boss; L21-30 two DIFFERENT mediums simultaneously, surge gap, end boss; L31-40 two "
     "mediums, surge gap, a transitional boss, then the end boss IMMEDIATELY (0s) on its "
     "death. Bosses never co-exist. The surge gap is rand(10,15)s and floods the arena with "
     "a DENSE wave of monsters (alive cap 28). Medium/transitional kinds reuse existing AIs "
     "plus two NEW boss AIs (warden, harrier) that are never tier end bosses, so validator "
     "C6's locked order is untouched."),

    # Leveling bullet -> account level cap 40 -> 50. (Rewrite-in-place; text matches the
    # INSERTS version so the insert step then skips it. find anchors on the OLD text.)
    ("xpToNext grows geometrically (cap 40)",
     "LEVELING + SCALING: XP from kills (small, accrued) + level-clear/victory "
     "(banked); xpToNext grows geometrically and steepens past L15 (account level "
     "cap 50). applyLevelScaling() at run start raises base damage, fire rate / "
     "attack speed, max HP, and the held-RMB 'second basic attack' DURATION (its "
     "fuel/energy/power pool starts short and grows with level). Drop rate also "
     "scales up with level (gate only)."),

    # PARAMETERS bullet -> add the distance unit (px) + attack RANGE. (Rewrite-in-place.)
    ("move speed / shield / regen / lifesteal)",
     "CHARACTER PARAMETERS screen (STATE.PARAMETERS): a PARAMETERS button on the MENU "
     "(top-right, just below the profile bar's COINS line) opens a menu-side version of "
     "the same readout, with three class tabs. Since there is no live player on the "
     "menu, values come from scaledStatsFor(charKey, level) - a pure mirror of "
     "applyLevelScaling - and each numeric row also shows its GAIN PER ACCOUNT LEVEL as "
     "a (+x) token (computed by diffing the helper at level vs level+1; '(+0)' for "
     "stats that don't scale, like move speed / shield / regen / lifesteal / range). It "
     "also manifests the DISTANCE UNIT (px = game pixels of the 480x270 world): MOVE "
     "SPEED in px/s, a RANGE row for LMB (projectile reach) and RMB (reach + scope) in "
     "px, and a bottom legend."),
]

# INSERTS: bullets added immediately before the paragraph whose text contains the
# heading key. (before_heading, [bullet_text, ...])
INSERTS = [
    ("6. RMB secondaries", [
        "META-PROGRESSION (persistent): one shared account profile in localStorage "
        "(shooter_profile) - level, XP, gold, owned skills + assistants, the X/C/B/V "
        "loadout, and equipped assistants. XP earned on ANY character credits the one "
        "profile, so all characters level together. loadProfile() guards a non-string "
        "(headless stubs localStorage as a Proxy).",

        "LEVELING + SCALING: XP from kills (small, accrued) + level-clear/victory "
        "(banked); xpToNext grows geometrically and steepens past L15 (account level "
        "cap 50). applyLevelScaling() at run start raises base damage, fire rate / "
        "attack speed, max HP, and the held-RMB 'second basic attack' DURATION (its "
        "fuel/energy/power pool starts short and grows with level). Drop rate also "
        "scales up with level (gate only).",

        "SHOP (STATE.SHOP, the ARMORY): spend Gold on tiered skills (cheap/small -> "
        "expensive/big) and on assistants - which you BUY then UPGRADE through 10 levels "
        "(the gold cost rises each level); purchases persist. LOADOUT screen: F is locked "
        "to HEAL, X/C/B/V unlock +1 per 5 levels and are filled from owned skills. The "
        "GUIDE button now sits top-LEFT; the persistent LV/XP bar shows in the top-RIGHT "
        "corner with the COINS total (a specific number) on its own line just below the "
        "XP bar.",

        "CHARACTER STATS window (Tab): during play, Tab opens a screen-space panel of "
        "the character's LIVE, level-scaled basic parameters - HP, shield, move speed, "
        "primary (LMB) and secondary (RMB) attack damage + fire cadence, the "
        "FUEL/ENERGY/POWER pool capacity (= second-attack capacity) and its per-second "
        "regen, the overall damage multiplier, Soldier lifesteal, and account level. It "
        "freezes the tick like pause (update() skips updatePlaying while showStats); "
        "Tab/Esc/click closes it. HP-regen is shown only for the Scout (0.5%/s) - the "
        "others read '-', because there is no universal HP-per-second regen.",

        "RMB SECONDARY POOLS (uniform across classes): all three resources (Syed energy, "
        "Benjamin fuel, Xu Yihui power) now self-regen ALWAYS - even while RMB is held - and "
        "all three get a fixed 1-second no-fire LOCKOUT the instant the pool empties "
        "(player.rmbLockT). The pay-per-use gate means the secondary deals ZERO damage while "
        "the pool can't cover a use; the lockout (not a regen pause) is what stops near-empty "
        "trickle-fire. applyLevelScaling grows the pool CAPACITY with account level.",

        "ASSISTANT STAT READOUTS: the SHOP ASSISTANTS tab (taller cards) and the LOADOUT "
        "screen both show each ally's level-scaled basic stats (HP / damage / range / rate of "
        "fire, or DPS/blast for the bomber/poison) via assistantStatLines(key), which builds a "
        "throwaway createAssistant so the numbers reflect the current profile.assistantLevels "
        "upgrade level - upgrading an ally in the shop visibly raises them.",

        "5-TIER SKILLS: every skill (incl. HEAL) now upgrades across 5 tiers in the SHOP - "
        "each tier gives a stronger effect AND a shorter cooldown, with a non-linear gold cost "
        "(cheap early, steep late). Each SKILLS[id] carries a t:[...] array of per-tier knobs; "
        "the active tier lives in profile.skillLevels[id], and triggerSkill reads skillStat(id) "
        "for every magnitude/duration while the dispatch + HUD read skillCdFor(id). Balance "
        "ramps weak->strong (tier 1 is weaker than the old single value, tier ~3 about matches "
        "it, tier 5 exceeds it); already-owned skills start at tier 1. The SHOP skills tab shows "
        "5 tier pips per card + a hover detail panel with the current effect and the next tier's "
        "effect in parens.",

        "ASSISTANT EVOLUTIONS + RESUMMON: each ally's upgrade level also unlocks behavioral "
        "evolutions (L1-3 base / L4-6 secondary / L7-10 ultimate) layered on the AI - e.g. the "
        "drone fires an AoE missile every 3rd shot at L4-6 and gains a periodic x2-firerate "
        "berserk at L7-10; brute slam+enrage, nunchaku whirlwind+flurry, bomber cluster+carpet, "
        "poison spore-nova+plague. RESUMMON is a new universal loadout skill that revives a dead "
        "equipped assistant near you (no-op, no cooldown spent, when none are dead; higher tiers "
        "grant brief invuln and can revive two).",

        "CHARACTER PARAMETERS screen (STATE.PARAMETERS): a PARAMETERS button on the MENU "
        "(top-right, just below the profile bar's COINS line) opens a menu-side version of "
        "the same readout, with three class tabs. Since there is no live player on the "
        "menu, values come from scaledStatsFor(charKey, level) - a pure mirror of "
        "applyLevelScaling - and each numeric row also shows its GAIN PER ACCOUNT LEVEL as "
        "a (+x) token (computed by diffing the helper at level vs level+1; '(+0)' for "
        "stats that don't scale, like move speed / shield / regen / lifesteal / range). It "
        "also manifests the DISTANCE UNIT (px = game pixels of the 480x270 world): MOVE "
        "SPEED in px/s, a RANGE row for LMB (projectile reach) and RMB (reach + scope) in "
        "px, and a bottom legend.",
    ]),
    ("10. Allies", [
        "FOUR-PHASE difficulty curve (difficultyRamp / enemyScale / bossHpScale / "
        "bossSpdScale): scaling is gentle through L1-15, moderate L16-24, sharp L25-30, then a "
        "STEEP L31-40 spike so the late game depends on meta-progression upgrades. The enemy "
        "HP/speed ramp's per-level step is capped at 0.2, so most of the L31-40 spike is carried "
        "by boss HP/speed plus enemyDmgScale(idx) (1.0 through L30, up to ~2.2 at L40, applied to "
        "enemy contact/bullet/gas damage). Two new regular enemies join the L11-40 spawn mixes: "
        "weaver (evasive bullet-dodger) and spitter (drops telegraphed poison pools). DROP "
        "variety is also level-gated in killEnemy - basic & sparse in L1-15 "
        "(health/resource/rapid/shield), widening at L16 (pierce/multi/damage/speed/magnet) and "
        "L25 (full pool) - while the static drop arrays still sum to 1.0 (validator C5) and the "
        "runtime filter renormalizes.",

        "Reusable boss abilities (escalate by levelIdx, no-op for clones): bossInvuln "
        "(1.2s damage-absorb shield), bossSummonExploders, bossSpawnClone, bossGasBomb, "
        "bossHeal. The gas pool (bossGasBomb) spawns on the player but now ARMS for 0.5s "
        "- a pulsing yellow telegraph ring that deals NO damage - so you can step out "
        "before it goes live (no 100%-hit instant chip).",

        "STAGED boss encounters by 10-level block (bossPhase state machine): L1-10 end "
        "boss only; L11-20 one medium boss at 50% progress, then a surge gap, then the end "
        "boss; L21-30 two DIFFERENT mediums simultaneously, surge gap, end boss; L31-40 two "
        "mediums, surge gap, a transitional boss, then the end boss IMMEDIATELY (0s) on its "
        "death. Bosses never co-exist. The surge gap is rand(10,15)s and floods the arena with "
        "a DENSE wave of monsters (alive cap 28). Medium/transitional kinds reuse existing AIs "
        "plus two NEW boss AIs (warden, harrier) that are never tier end bosses, so validator "
        "C6's locked order is untouched.",
    ]),
]


def set_paragraph_text(p, text):
    """Replace a paragraph's text while keeping its paragraph style. Collapses to a
    single run, preserving the first run's character formatting where possible."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def main():
    if not os.path.exists(DOCX):
        print("ERROR: not found: " + DOCX, file=sys.stderr)
        return 2
    doc = docx.Document(DOCX)
    changed = 0

    # --- Replacements ---
    for find, new in REPLACEMENTS:
        hit = next((p for p in doc.paragraphs if find in p.text), None)
        if hit is not None:
            set_paragraph_text(hit, new)
            changed += 1
            print("[edit] " + find[:60])
        else:
            # Either already updated (idempotent re-run) or the doc drifted.
            already = any(new[:40] in p.text for p in doc.paragraphs)
            print(("[ok  ] already current: " if already else "[MISS] not found: ")
                  + find[:60])

    # --- Inserts (idempotent: skip a bullet whose lead text already exists) ---
    for heading, bullets in INSERTS:
        target = next((p for p in doc.paragraphs if heading in p.text), None)
        if target is None:
            print("[MISS] heading not found: " + heading)
            continue
        existing = [p.text for p in doc.paragraphs]
        for b in bullets:
            if any(b[:45] in e for e in existing):
                print("[ok  ] bullet already present: " + b[:50])
                continue
            target.insert_paragraph_before(b, style="List Bullet")
            changed += 1
            print("[add ] bullet before '%s': %s" % (heading, b[:50]))

    doc.save(DOCX)
    print("\nSaved %s (%d change(s) this run)." % (os.path.basename(DOCX), changed))
    return 0


if __name__ == "__main__":
    sys.exit(main())
